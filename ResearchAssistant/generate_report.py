import json
import os
import datetime
from typing import Type
import requests
from bs4 import BeautifulSoup
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from docx import Document
from dotenv import load_dotenv
from langchain.agents import initialize_agent, Tool, AgentType
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import PyPDFLoader
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.memory import ConversationSummaryBufferMemory
from langchain.prompts import MessagesPlaceholder
from langchain.schema import SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.tools import BaseTool
from langchain.vectorstores import FAISS
from pydantic import BaseModel, Field
from rest_framework import status
from rest_framework.response import Response
from .text_to_report import create_pdf, create_docx, create_pptx, create_excel
import re
import logging
from django.conf import settings


# Loading the environment variables from .env file
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
browserless_api_key = os.getenv("BROWSERLESS_API_KEY")
serper_api_key = os.getenv("SERP_API_KEY")
huggingfacehub_api_token = os.getenv("HUGGINGFACEHUB_API_TOKEN")

# Declare the Constants
# API_BASE_URL = 'http://127.0.0.1:8000'
API_BASE_URL = 'http://123.201.192.65:8282'
API_ENDPOINT = '/api/inputs-report-generation/'
Research_History_API = f"{API_BASE_URL}/api/research-history/"


class CreateDocument:
    def __init__(self, content):
        self.content = content


def process_uploaded_file(file_path):
    # Check the file type using file extension and return file details

    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        process_pdf(file_path)

    elif file_extension == '.txt':
        process_text(file_path)

    elif file_extension == '.docx':
        process_docx(file_path)

    else:
        return JsonResponse({'message': 'Please upload pdf or text or word document'}, status=400)


# Function to process PDF file
def process_pdf(file_path):
    # print("processing PDF file")
    loader = PyPDFLoader(file_path)
    pages = loader.load_and_split()
    # Define chunk size, overlap and separators
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=64,
        separators=['\n\n', '\n', '(?=>\. )', ' ', '']
    )
    docs = text_splitter.split_documents(pages)

    # Embeddings
    embeddings = HuggingFaceEmbeddings()
    # print("PDF file embedding done")
    # Create the vectorized db
    # Vectorstore: https://python.langchain.com/en/latest/modules/indexes/vectorstores.html

    db = FAISS.from_documents(docs, embeddings)
    db.save_local("faiss_index")
    new_db = FAISS.load_local("faiss_index", embeddings)
    # print("PDF file-embedding-indexing done- saved into VECTOR DB")


# Function to process text file
def process_text(file_path):
    # print("processing TEXT file")
    with open(file_path, "r") as file:
        text_data = file.read()

    # Define chunk size, overlap, and separators
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=64,
        separators=['\n\n', '\n', '(?=>\. )', ' ', '']
    )

    # Create a Document instance with text_data as content
    document = CreateDocument(text_data)

    # Embeddings
    embeddings = HuggingFaceEmbeddings()
    # print("TEXT file embedding done")

    # Create the vectorized db
    # Vectorstore: https://python.langchain.com/en/latest/modules/indexes/vectorstores.html
    db = FAISS.from_documents([document], embeddings)
    db.save_local("faiss_index")
    new_db = FAISS.load_local("faiss_index", embeddings)
    # print("TEXT file-embedding-indexing done- saved into VECTOR DB")


# Function to process docx file
def process_docx(file_path):
    # print("processing WORD DOC file")
    doc = Document(file_path)
    text_data = ""
    for para in doc.paragraphs:
        text_data += para.text + "\n"

    # Define chunk size, overlap, and separators for text splitting

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=64,
        separators=['\n\n', '\n', '(?=>\. )', ' ', '']
    )

    # Create a Document instance with text_data as content
    document = CreateDocument(text_data)

    # Embeddings
    embeddings = HuggingFaceEmbeddings()
    # print("WORD DOC file embedding done")
    # Create the vectorized db
    # Vectorstore: https://python.langchain.com/en/latest/modules/indexes/vectorstores.html
    db = FAISS.from_documents([document], embeddings)
    db.save_local("faiss_index")
    new_db = FAISS.load_local("faiss_index", embeddings)
    # print("WORD DOC file-embedding-indexing done-saved into VECTOR DB")


def research_report_generation(request):
    # Extract user ID from the logged-in user
    user_id = request.user.id
    # Construct the full API URL
    api_url = f"{API_BASE_URL}{API_ENDPOINT}{user_id}/"
    # print('GET API Created :: ', api_url)

    try:

        # Make a GET request to the API
        response = requests.get(api_url)
        # print('GET API Called')

        # Check if the request was successful (status code 200)
        if response.status_code == 200:
            # Parse the JSON response
            data = response.json()
            # print("Request from User:", data['user'])
            # print("User Inputs Fetched Successfully")
            # If an uploaded file exists, process it
            research_goal = data.get('research_goal')
            research_objective = data.get('research_objective')
            research_parameters = data.get('research_parameters')
            additional_text_input = data.get('additional_text_input')
            report_format = data.get('report_format')
            uploaded_file_path = data.get('uploaded_file')
            result = None
            if uploaded_file_path:
                # print("User has uploaded the file")
                # print("Uploaded File Path is:", uploaded_file_path)
                process_uploaded_file(uploaded_file_path)
                # print("Uploaded file processed- Calling Research Agent for getting Report")
                result = research_agent(research_goal, research_objective, research_parameters, additional_text_input)
                research_topic, sources = extract_research_topic_and_domains(result, user_id)
                create_research_history(request, report_format, result, research_topic, sources)

            else:
                # print("User has not uploaded the file- Calling Research Agent for Report")
                result = research_agent(research_goal, research_objective, research_parameters, additional_text_input)
                research_topic, sources = extract_research_topic_and_domains(result, user_id)
                create_research_history(request, report_format, result, research_topic, sources)

            return result, sources

        else:
            # Print an error message if the request was not successful
            print(f"Error: Unable to fetch data. Status code {response.status_code}")

    except requests.exceptions.RequestException as e:
        # Handle any exceptions that occurred during the request
        print(f"Error: {e}")

    # Return None if there's an error or no result from research_agent
    return None


# Code of AI Engine starts below *********************************************************************

# 1. Tool for search

def search(query):
    # print("Searching the internet for urls")
    url = "https://google.serper.dev/search"

    payload = json.dumps({
        "q": query
    })

    headers = {
        'X-API-KEY': serper_api_key,
        'Content-Type': 'application/json'
    }

    try:
        # Make the POST request
        response = requests.post(url, headers=headers, data=payload)

        # Check if the request was successful (status code 200)
        if response.status_code == 200:
            # Parse the JSON response
            search_results = response.json()

            # Extract URLs from the search results
            urls = extract_urls(search_results)

            return urls
        else:
            print(f"Error: Unable to fetch search results. Status code {response.status_code}")

    except requests.exceptions.RequestException as e:
        # Handle any exceptions that occurred during the request
        print(f"Error: {e}")

    # Return an empty list if there's an error or no URLs are found
    return []


def extract_urls(search_results):
    try:
        # Assuming search_results is in JSON format
        urls = [item.get('link') for item in search_results.get('items', [])]
        return urls
    except Exception as e:
        print(f"Error extracting URLs from search results: {e}")
        return []


# 2. Tool for scraping
def scrape_website(objective: str, url: str):
    # scrape website, and also will summarize the content based on objective if the content is too large
    # objective is the original objective & task that user give to the agent, url is the url of the website to be scraped

    # print("Website Scraping Started **")
    # Define the headers for the request
    headers = {
        'Cache-Control': 'no-cache',
        'Content-Type': 'application/json',
    }

    # Define the data to be sent in the request
    data = {
        "url": url
    }

    # Convert Python object to JSON string
    # Get URL domains
    data_json = json.dumps(data)

    # Send the POST request
    post_url = f"https://chrome.browserless.io/content?token={browserless_api_key}"
    response = requests.post(post_url, headers=headers, data=data_json)

    # Check the response status code
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, "html.parser")
        # print("Getting TEXT from scraped websites")
        text = soup.get_text()
        text_splitter = RecursiveCharacterTextSplitter(
            separators=["\n\n", "\n"], chunk_size=10000, chunk_overlap=1000)
        docs = text_splitter.create_documents(text)
        # Embeddings
        # print("Doing website content embedding")
        embeddings = HuggingFaceEmbeddings()
        # embeddings = OpenAIEmbeddings()

        # Create the vectorized db
        # Vectorstore: https://python.langchain.com/en/latest/modules/indexes/vectorstores.html
        db = FAISS.from_documents(docs, embeddings)
        db.save_local("faiss_index")
        new_db = FAISS.load_local("faiss_index", embeddings)
        # print("Saved scraped website content in VECTOR DB")
    else:
        print(f"HTTP request failed with status code {response.status_code}")


def perform_similarity_search(new_db, objective: str):
    # perform similarity search with the saved new_db
    relevant_content = new_db.similarity_search_with_score(objective)
    # print("Got the relevant_content from VECTOR DB")
    return relevant_content


def process_websites(objective: str, urls: list):
    # Iterate over the list of URLs
    # print("Processing websites one by one")
    for url in urls:
        # Scrape the website and save its data into a vector database
        scrape_website(objective, url)

    # Perform similarity search on the aggregated vector database
    new_db = FAISS.load_local("faiss_index", HuggingFaceEmbeddings())  # Load the vector database
    relevant_content = perform_similarity_search(new_db, objective)

    return relevant_content


class ScrapeWebsiteInput(BaseModel):
    """Inputs for scrape_website"""
    objective: str = Field(
        description="The objective & task that users give to the agent")
    url: str = Field(description="The url of the website to be scraped")


class ScrapeWebsiteTool(BaseTool):
    # process_websites  TRY
    # name = "scrape_website"
    name = "process_websites"
    description = "useful when you need to get data from a website url, passing both url and objective to the function; DO NOT make up any url, the url should only be from the search results"
    args_schema: Type[BaseModel] = ScrapeWebsiteInput

    def _run(self, objective: str, urls: list):
        # process_websites  TRY
        return process_websites(objective, urls)

    def _arun(self, url: str):
        raise NotImplementedError("error here")


# 3. Create langchain agent with the tools above
tools = [
    Tool(
        name="Search",
        func=search,
        description="useful for when you need to answer questions about current events, data. You should ask targeted questions"
    ),
    ScrapeWebsiteTool(),
]

system_message = SystemMessage(
    content="""You are a world class researcher, who can do detailed research on any topic and produce facts based results; 
        you do not make things up, you will try as hard as possible to gather facts & data to back up the research

        Please make sure you complete the objective above to write a high quality research report using the events and data 
        that can be useful to market analysts with industry analysis, competitor analysis, consumer trends etc., with the following rules:
        1/ Start with appropriate Topic heading for the research report and give a detailed research report with relevant sources
        2/ Include necessary infographics or visualizations that describe the market trends, challenges and opportunities
        3/ You should do enough research to gather as much information as possible about the objective
        4/ If there are url of relevant links & articles, you will scrape it to gather more information
        5/ After scraping & search, you should think "is there any new things i should search & scraping based on the data I collected to increase research quality?" If answer is yes, continue; But don't do this more than 3 iterations
        6/ You should not make things up, you should only write facts & data that you have gathered
        7/ In the final output, You should include all reference data & links to back up your research;
        """
)

agent_kwargs = {
    "extra_prompt_messages": [MessagesPlaceholder(variable_name="memory")],
    "system_message": system_message,
}

llm = ChatOpenAI(temperature=0, model="gpt-3.5-turbo-16k-0613" )
memory = ConversationSummaryBufferMemory(
    memory_key="memory", return_messages=True, llm=llm, max_token_limit=10000)

agent = initialize_agent(
    tools,
    llm,
    agent=AgentType.OPENAI_FUNCTIONS,
    verbose=False,
    agent_kwargs=agent_kwargs,
    memory=memory,
)


def research_agent(research_goal, research_objective, research_parameters, additional_text_input=None):
    # Check if additional_text_input is provided
    if additional_text_input is not None:
        query = f"{research_goal} {research_objective} {research_parameters} {additional_text_input}"
    else:
        query = f"{research_goal} {research_objective} {research_parameters}"

    content = agent({"input": query})
    actual_report = content['output']
    # print('Generated Report ****' * 5)
    return actual_report


# ------------------------------Research History Code---------------------------------------------
# Extract research topic and domains function
def extract_research_topic_and_domains(report_content, user_id):
    current_datetime = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    research_topic_match = re.search(r"Research Report: (.+?)\n", report_content)
    research_topic = research_topic_match.group(1) if research_topic_match else f'SystemicLogic Intellisense Report {user_id}_{current_datetime}'
    references_match = re.findall(r"https?://(?:www\.)?([a-zA-Z0-9.-]+)", report_content)
    sources = list(set(match for match in references_match))

    return research_topic, sources


def create_research_history(request, report_format, actual_content, research_topic, sources):
    try:
        # Check if report_format and actual_content are available
        if report_format is None or actual_content is None:
            raise ValueError("report_format or actual_content is not set.")

        # research_topic, sources = extract_research_topic_and_domains(actual_content)
        user_id = request.user.id

        if report_format == 'pdf':
            content, files = create_pdf(user_id, actual_content, research_topic)
        elif report_format == 'docx':
            content, files = create_docx(user_id, actual_content, research_topic)
        elif report_format == 'pptx':
            content, files = create_pptx(user_id, actual_content, research_topic)
        elif report_format == 'xlsx':
            content, files = create_excel(user_id, actual_content, research_topic)
        else:
            raise ValueError(f"Unsupported report format: {report_format}")

        # POST data to the endpoint
        url = Research_History_API

        data = {
            "research_title": research_topic,
            "sources": ', '.join(sources),
            "user": user_id
        }

        # Upload file for generated_report field
        response = requests.post(url, data=data, files=files)

        if response.status_code == status.HTTP_201_CREATED:
            # Log success or other relevant information
            logging.info("Data successfully posted to ResearchHistory table.")
            return Response({"message": "Function executed successfully"}, status=status.HTTP_200_OK)
        else:
            error_message = f"Failed to post data. Status code: {response.status_code}, Response text: {response.text}"
            logging.error(error_message)
            return Response({"error": error_message}, status=response.status_code)

    except Exception as e:
        # Log any unexpected exceptions
        logging.exception("An unexpected error occurred.")
        return Response({"error": f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
